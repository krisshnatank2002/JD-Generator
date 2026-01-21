# jd_generator.py

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage
import streamlit as st

# =====================================================
# FONT SIZES
# =====================================================
TITLE_FONT_SIZE = Pt(13)
HEADING_FONT_SIZE = Pt(11)
BODY_FONT_SIZE = Pt(09)

# =====================================================
# LOAD GROQ API KEY
# =====================================================
try:
    GROQ_API_KEY = st.secrets["GROQ_API_KEY"]
except KeyError:
    raise RuntimeError("GROQ_API_KEY not found in Streamlit Secrets")

llm = ChatGroq(
    model="llama-3.3-70b-versatile",
    temperature=0,
    api_key=GROQ_API_KEY
)

# =====================================================
# TITLE CASE HELPER
# =====================================================
def to_title_case(title: str) -> str:
    if not title:
        return title
    words = title.split()
    return " ".join(w.upper() if w.lower() == "ai" else w.capitalize() for w in words)
# =====================================================
# CONSTANT COMPANY DESCRIPTION (LOCKED)
# =====================================================
ABOUT_WOGOM_TEXT = """WOGOM is a B2B Commerce and Retail Enablement Platform, empowering 6,000+ retailers
and 450+ sellers across India with better Products, Pricing, Credit, and Growth
Opportunities.
Our goal is to build a connected ecosystem where technology, capital, and commerce converge
to help Indian retailers scale with confidence."""
# =====================================================
# JD HEADINGS
# =====================================================
HEADINGS = {
    "Reporting To",
    "About WOGOM",
    "Role Overview",
    "What You’ll Do?",
    "What You'll Do?",
    "Who’ll Succeed in this Role?",
    "Who'll Succeed in this Role?",
    "Must-Have Skills",
    "Preferred Skills",
    "Hiring Priority",
}

# =====================================================
# DOCX HELPERS
# =====================================================
def remove_numbering(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)

def add_job_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = TITLE_FONT_SIZE

def add_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = HEADING_FONT_SIZE
    remove_numbering(p)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = BODY_FONT_SIZE
    remove_numbering(p)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = BODY_FONT_SIZE

# def add_inline_skills(doc, skills):
#     """
#     skills: list[str]
#     Converts bullet skills into one compact line
#     """
#     line = ", ".join(skills)
#     add_paragraph(doc, line)

# =====================================================
# CTC & JOINING BLOCK (MANDATORY)
# =====================================================
def add_ctc_and_joining(doc, row):
    salary = ""
    urgency = ""

    if "Salary range (optional)" in row and str(row["Salary range (optional)"]).strip():
        salary = str(row["Salary range (optional)"]).strip()

    if "How urgent is this hire?" in row and str(row["How urgent is this hire?"]).strip():
        urgency = str(row["How urgent is this hire?"]).strip()

    # ✅ ALWAYS SHOW THIS SECTION
    add_heading(doc, "Compensation & Joining")

    if salary:
        add_paragraph(doc, f"CTC: {salary}")
    else:
        add_paragraph(doc, "CTC: As per company standards")

    if urgency:
        add_paragraph(doc, f"Joining: {urgency}")
    else:
        add_paragraph(doc, "Joining: As per mutual availability")

# =====================================================
# HEADER BLOCK
# =====================================================
def build_header_block(row):
    parts = []

    for key in row.index:
        if "location" in key.lower():
            parts.append(row[key])
        if "employment" in key.lower():
            parts.append(row[key])
        if "work mode" in key.lower() or "workmode" in key.lower():
            parts.append(row[key])

    travel = row.get("Does this role require travel?", "").strip()
    if travel:
        parts.append(f"{travel} travel")

    return " | ".join([p for p in parts if p])

# =====================================================
# CLARIFICATION SANITIZER
# =====================================================
def sanitize_clarifications(clarifications):
    """
    Keeps your OLD clarification structure.
    If user selected 'None of the above' / 'Not Applicable',
    that clarification is REMOVED so it does NOT affect the JD.
    """
    sanitized = {}
    for question, answer in clarifications.items():
        if isinstance(answer, str):
            if answer.strip().lower() in ["not applicable", "none of the above"]:
                continue
            sanitized[question] = answer
    return sanitized

# =====================================================
# CORE JD GENERATION
# =====================================================
def generate_ranked_jd(row, clarifications=None):
    """
    Clarifications are MANDATORY if provided.
    They OVERRIDE assumptions and must influence:
    - Responsibilities
    - Requirements
    - Skills
    - Seniority
 
    If clarification == 'Not Applicable', it is ignored.
    """

    clarifications = clarifications or {}
    clarifications = sanitize_clarifications(clarifications)

    # 🔹 Detect Job Title column dynamically
    job_title_col = None
    for k in row.index:
        if "job" in k.lower() and "title" in k.lower():
            job_title_col = k
            break

    job_title = row.get(job_title_col, "")

    # 🔹 Override if clarified
    for q, a in clarifications.items():
        if "job title" in q.lower():
            job_title = a
            break

    job_title = to_title_case(job_title)

    clarification_block = ""
    if clarifications:
        clarification_block = """
IMPORTANT – MANDATORY OVERRIDE RULES:
 
You MUST strictly incorporate ALL clarifications below.
These clarifications OVERRIDE assumptions from input data.
They MUST directly affect responsibilities, scope, skills, and expectations.
 
HIRING MANAGER CLARIFICATIONS:
"""
        for q, a in clarifications.items():
            clarification_block += f"- {q}: {a}\n"

    prompt = f"""
{clarification_block}

STRICT FORMATTING RULES(NON-NEGOTIABLE):
GENERAL TONE:
- Write like a hiring manager, not HR
- Focus on execution, ownership, and outcomes
- Avoid generic phrases (e.g. "strong communication", "team player", "dynamic environment")
- Remove repetition across sections
- Be confident, direct, and concise

Role Title section:
- Output ONLY: "Role Title" heading followed by the job title.
- DO NOT include location, travel, or meta info.
 
Skills Sections:
- Bullet points ONLY
- Each skill = keyword or short phrase (no sentences)
- No repetition between Must-Have and Preferred
- Keep skills concise and scannable
If a clarification is not provided, DO NOT infer or assume details for it.
 
=====================
REQUIRED STRUCTURE
=====================
 
Role Title
<Job Title Only>

About WOGOM
<Use the provided company description verbatim>

Role Overview
<1 3-4 line clear paragraph explaining the role's purpose and impact.Direct, outcome-focused. No fluff.>
 
What You'll Do?
Start with a short paragraph (2–3 lines) summarizing the role’s execution focus.

Then list 8–10 key responsibilities.
Rules:
- Use bullet points only
- Each bullet starts with an action verb
- little explanations or filler
 

Who’ll Succeed in this Role?
List skills and expectations using bullet points only.

Structure:
- Start directly with bullets (no paragraph)
- Each bullet follows this format:
  Skill / Requirement – short 4–6 word description

After listing must-have items, add a line:
"Preferred Skills"

Under Preferred Skills:
- Continue bullet points
- 2–4 items only
- Same format
- Do NOT repeat must-have items
Rules:
- Keep this section concise and scannable
- No motivational language
- No personality adjectives
- No repetition from “What You’ll Do?”


 
=====================
INPUT DATA
=====================
 
Job Title: {row.get('Job Title','')}
 
Reporting To: {row.get('Reporting To','')}
 
Role Overview:
{row.get('Role Overview','')}
 
Education: {row.get('Minimum education required','')}
Experience: {row.get('Minimum experience required','')}
Core Responsibility: {row.get('What is the single core responsibility of this role?','')}
 
Key Responsibilities:
{row.get('Key Responsibilities','')}
 
Top Skills:
{row.get('Top 3 skills this role MUST have','')}
 
Other Skills:
{row.get('other skills','')}
 
Growth Opportunities:
{row.get('Growth opportunities in this role','')}
 
Company Context:
{row.get('Role Context','')}
"""
    response = llm.invoke([HumanMessage(content=prompt)])
    return response.content.strip()

# =====================================================
# WRITE JD TO DOCX
# =====================================================
def write_jd_to_docx(jd_text, row):
    doc = Document()

    BULLET_SECTIONS = {
        "Responsibilities",
        "Requirements",
        "Must-Have Skills",
        "Preferred Skills",
    }
    # Job title
    add_job_title(doc, row["__job_title__"])

    # Meta line
    meta = build_header_block(row)
    if meta:
        add_paragraph(doc, meta)

    # Parse JD
    lines = [l.strip() for l in jd_text.split("\n") if l.strip()]
    current_section = None
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip Role Title section from LLM
        if line == "Role Title":
            i += 2
            continue

        if line in HEADINGS:
            add_heading(doc, line)
            current_section = line
            i += 1
            continue

        if current_section in {"Must-Have Skills", "Preferred Skills"}:
            skills = []

            while (
                i < len(lines)
                and lines[i] not in HEADINGS
                and not lines[i].endswith("?")
            ):
                skills.append(lines[i].lstrip("•-* ").strip())
                i += 1

            if skills:
                add_inline_skills(doc, skills)

            continue

        # Handle bullets everywhere except title
        if line.startswith(("•", "-", "*")):
            clean = line.lstrip("•-* ").strip()
            add_bullet(doc, clean)
            i += 1
        continue



        add_paragraph(doc, line)
        i += 1
    add_ctc_and_joining(doc, row)
    return doc








